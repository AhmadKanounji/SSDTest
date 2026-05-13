from datetime import datetime
from docx import Document
from copy import deepcopy


TEMPLATE_PATH = "RN_EG_CBE_Billing_Template.docx"
OUTPUT_PATH = "Release_Note_Output.docx"


def clear_table_data_rows(table):
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[2]._tr)

    if len(table.rows) > 1:
        for cell in table.rows[1].cells:
            cell.text = ""


def add_row(table, values):
    # If row 1 is still empty, use it directly
    target_row = table.rows[1]

    is_empty = all(
        cell.text.strip() == ""
        for cell in target_row.cells
    )

    if not is_empty:
        new_tr = deepcopy(target_row._tr)
        table._tbl.append(new_tr)
        target_row = table.rows[-1]

    for i, value in enumerate(values):
        target_row.cells[i].text = str(value or "")

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, str(value))


def build_release_note_docx(release_version, release_author, fixed_rows, open_rows):
    doc = Document(TEMPLATE_PATH)

    today = datetime.now().strftime("%d/%m/%Y")

    component = fixed_rows[0]["platform"] if fixed_rows else ""

    replace_placeholders(doc, {
        "{{Component}}": component,
        "{{COMPONENT}}": component,
        "{{RELEASE_VERSION}}": release_version,
        "{{GENERATION_DATE}}": today,
    })

    # Revision History = table 2
    revision_table = doc.tables[2]
    clear_table_data_rows(revision_table)
    add_row(revision_table, [
        release_version,
        today,
        release_author,
        "Creation of document"
    ])

    # Fixed Tickets = table 5
    fixed_table = doc.tables[5]
    clear_table_data_rows(fixed_table)

    for row in fixed_rows:
        add_row(fixed_table, [
            row.get("issue_key", ""),
            row.get("external_id", ""),
            row.get("summary", ""),
            row.get("platform", ""),
            row.get("severity", ""),
            row.get("fixed_version", ""),
        ])

    # Open Tickets = table 6
    open_table = doc.tables[6]
    clear_table_data_rows(open_table)

    for row in open_rows:
        add_row(open_table, [
            row.get("issue_key", ""),
            row.get("external_id", ""),
            row.get("summary", ""),
            row.get("platform", ""),
            row.get("severity", ""),
        ])

    doc.save(OUTPUT_PATH)

    return OUTPUT_PATH
