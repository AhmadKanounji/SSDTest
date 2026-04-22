import os
import re
import tempfile
import html
from datetime import datetime
from zoneinfo import ZoneInfo
from pathlib import Path

import requests
from requests.auth import HTTPBasicAuth
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- ENV ----------
ATLASSIAN_DOMAIN = os.environ["ATLASSIAN_DOMAIN"].strip()
EMAIL = os.environ["ATLASSIAN_EMAIL"].strip()
API_TOKEN = os.environ["ATLASSIAN_API_TOKEN"].strip()
PROJECT_KEY = os.environ["PROJECT_KEY"].strip()

JIRA_BASE = f"https://{ATLASSIAN_DOMAIN}"
auth = HTTPBasicAuth(EMAIL, API_TOKEN)

# ---------- CONFIG ----------
TEMPLATE_PATH = os.environ.get("SSD_TEMPLATE_PATH", "ssd_template.docx")
OUTPUT_PATH = os.environ.get("SSD_OUTPUT_PATH", "SSD_Output.docx")
TZ = "Africa/Cairo"
PURPLE_HEX = "7030A0"


def jira_search(jql: str):
    url = f"{JIRA_BASE}/rest/api/3/search/jql"
    params = {
        "jql": jql,
        "maxResults": 200,
        "fields": "summary,description,issuetype,parent,attachment",
    }
    r = requests.get(url, params=params, auth=auth)
    r.raise_for_status()
    return r.json()["issues"]


def adf_to_text(adf):
    if not adf:
        return ""
    if isinstance(adf, str):
        return adf

    parts = []

    def walk(node):
        if isinstance(node, dict):
            node_type = node.get("type")
            if node_type == "text":
                parts.append(node.get("text", ""))
            elif node_type == "hardBreak":
                parts.append("\n")
            for child in node.get("content", []):
                walk(child)
            if node_type in ("paragraph", "heading", "listItem"):
                parts.append("\n")
        elif isinstance(node, list):
            for item in node:
                walk(item)

    walk(adf)
    text = "".join(parts)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_requirement_text(text: str) -> str:
    if not text:
        return ""
    cleaned = re.sub(r"\[(?:[^\]]+)\]+", "", text)
    cleaned = re.sub(r"\s+-\s+-\s+", " - ", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip(" -")


def first_image_attachment(issue):
    attachments = issue["fields"].get("attachment") or []
    for att in attachments:
        mime = (att.get("mimeType") or "").lower()
        filename = (att.get("filename") or "").lower()
        if mime.startswith("image/") or filename.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
            return att
    return None


def download_attachment(att):
    if not att or not att.get("content"):
        return None
    r = requests.get(att["content"], auth=auth)
    r.raise_for_status()
    suffix = Path(att.get("filename") or "image.bin").suffix or ".bin"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(r.content)
    tmp.close()
    return tmp.name


def extract_existing_revision_rows_from_confluence(page_html: str):
    match = re.search(
        r"<h1[^>]*>\s*Revision History\s*</h1>\s*(<table\b.*?</table>)",
        page_html or "",
        flags=re.DOTALL | re.IGNORECASE
    )
    if not match:
        return []

    table_html = match.group(1)
    rows = []
    tr_matches = re.findall(r"<tr[^>]*>(.*?)</tr>", table_html, flags=re.DOTALL | re.IGNORECASE)

    for tr in tr_matches:
        td_matches = re.findall(r"<td[^>]*>(.*?)</td>", tr, flags=re.DOTALL | re.IGNORECASE)
        if len(td_matches) != 4:
            continue
        cleaned = [
            html.unescape(
                re.sub(r"<[^>]+>", "", td).replace("&nbsp;", " ").strip()
            )
            for td in td_matches
        ]
        rows.append({
            "version": cleaned[0],
            "date": cleaned[1],
            "author": cleaned[2],
            "modification": cleaned[3],
        })
    return rows


def get_confluence_page(page_id):
    conf_base = f"https://{ATLASSIAN_DOMAIN}/wiki"
    url = f"{conf_base}/rest/api/content/{page_id}"
    params = {"expand": "body.storage,version"}
    r = requests.get(url, params=params, auth=auth)
    r.raise_for_status()
    return r.json()


def extract_use_case_sort_key(summary: str, fallback_key: str = ""):
    text = (summary or "").strip()
    normalized = text.lower()
    if normalized in ("exigences générales", "exigences generales"):
        return (0, [], normalized, fallback_key)

    patterns = [
        r"\bUC\s*([0-9]+(?:\.[0-9]+)*)\b",
        r"\bUse\s*Case\s*([0-9]+(?:\.[0-9]+)*)\b",
        r"^\s*([0-9]+(?:\.[0-9]+)*)\b",
    ]
    value = None
    for pattern in patterns:
        m = re.search(pattern, text, flags=re.IGNORECASE)
        if m:
            value = m.group(1)
            break

    if value is None:
        return (2, [999999], normalized, fallback_key)

    try:
        return (1, [int(p) for p in value.split(".")], normalized, fallback_key)
    except Exception:
        return (2, [999999], normalized, fallback_key)


def set_run_font(run, name="Arial", size=9, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

    if color is not None:
        run.font.color.rgb = color

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:cs"), name)


def add_cover_values(doc, version, date):
    replacements = {
        "Document Reference Number:": "",
        "Document Release Version:": version,
        "Document Release Date:": date,
    }
    for p in doc.paragraphs[:20]:
        txt = p.text.strip()
        for label, value in replacements.items():
            if txt.startswith(label):
                p.clear()
                run1 = p.add_run(label)
                set_run_font(run1, size=8, bold=True)
                if value:
                    run2 = p.add_run(" " + value)
                    set_run_font(run2, size=8)
                break


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, name="Arial", size=9, bold=bold, color=color)


def style_distribution_list_table(doc):
    if len(doc.tables) < 1:
        return

    table = doc.tables[0]
    table.style = "Table Grid"

    headers = ["Name", "Company"]
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
        for i, header in enumerate(headers):
            set_cell_background(table.rows[0].cells[i], PURPLE_HEX)
            set_cell_text(
                table.rows[0].cells[i],
                header,
                bold=True,
                color=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    for row in table.rows[1:]:
        for cell in row.cells:
            set_cell_text(cell, cell.text.strip(), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def fill_revision_history(doc, rows):
    if len(doc.tables) < 2:
        return

    table = doc.tables[1]
    table.style = "Table Grid"

    headers = ["Version", "Date", "Author", "Modification"]
    if len(table.rows) == 0:
        table.add_row()

    if len(table.rows[0].cells) >= 4:
        for i, header in enumerate(headers):
            set_cell_background(table.rows[0].cells[i], PURPLE_HEX)
            set_cell_text(
                table.rows[0].cells[i],
                header,
                bold=True,
                color=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row.get("version", ""))
        set_cell_text(cells[1], row.get("date", ""))
        set_cell_text(cells[2], row.get("author", ""))
        clean_mod = clean_requirement_text(row.get("modification", ""))
        set_cell_text(cells[3], clean_mod)


def fill_reference_documents_table(doc):
    if len(doc.tables) < 3:
        return

    table = doc.tables[2]
    table.style = "Table Grid"

    headers = ["Id", "Document or Meeting Name", "Release", "Date", "Reference"]
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 5:
        for i, header in enumerate(headers):
            set_cell_background(table.rows[0].cells[i], PURPLE_HEX)
            set_cell_text(
                table.rows[0].cells[i],
                header,
                bold=True,
                color=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )


def find_paragraph(doc, exact_text):
    for p in doc.paragraphs:
        if p.text.strip() == exact_text.strip():
            return p
    return None


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    if style:
        new_paragraph.style = style
    if text:
        run = new_paragraph.add_run(text)
        set_run_font(run, name="Arial", size=9)
    return new_paragraph


def insert_body_text_after(anchor_paragraph, text):
    current = anchor_paragraph
    for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
        current = insert_paragraph_after(current)
        current.alignment = WD_ALIGN_PARAGRAPH.LEFT
        current.paragraph_format.space_after = Pt(4)
        run = current.add_run(block)
        set_run_font(run, name="Arial", size=9)
    return current


def insert_requirement_title_after(anchor_paragraph, text):
    p = insert_paragraph_after(anchor_paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name="Arial", size=10, bold=True)
    return p


def insert_heading_2_after(anchor_paragraph, text):
    p = insert_paragraph_after(anchor_paragraph, style="Heading 2")
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name="Arial", size=10, bold=True)
    return p


def main():
    template = Document(TEMPLATE_PATH)

    page_id = os.environ.get("CONFLUENCE_PAGE_ID")
    existing_rows = []
    if page_id:
        page = get_confluence_page(page_id)
        existing_rows = extract_existing_revision_rows_from_confluence(
            page.get("body", {}).get("storage", {}).get("value", "")
        )
        existing_rows = list(reversed(existing_rows))

    today = datetime.now(ZoneInfo(TZ)).strftime("%d/%m/%Y")
    latest_version = existing_rows[0]["version"] if existing_rows else "0.1"

    add_cover_values(template, latest_version, today)
    style_distribution_list_table(template)
    fill_revision_history(template, existing_rows)
    fill_reference_documents_table(template)

    issues = jira_search(f'project = {PROJECT_KEY} AND issuetype in ("Use Case", Requirement)')
    use_cases = []
    reqs_by_uc = {}

    for issue in issues:
        issue_type = issue["fields"]["issuetype"]["name"]
        if issue_type == "Use Case":
            use_cases.append(issue)
        elif issue_type == "Requirement":
            parent = issue["fields"].get("parent")
            if parent and parent.get("key"):
                reqs_by_uc.setdefault(parent["key"], []).append(issue)

    use_cases = sorted(
        use_cases,
        key=lambda x: extract_use_case_sort_key(x["fields"].get("summary", ""), x.get("key", ""))
    )

    general = None
    regular = []
    for uc in use_cases:
        title = (uc["fields"].get("summary", "") or "").strip().lower()
        if title in ("exigences générales", "exigences generales"):
            general = uc
        else:
            regular.append(uc)

    # Inject Exigences Générales exactly under template headings
    if general:
        desc_anchor = find_paragraph(template, "3.1 Description")
        req_anchor = find_paragraph(template, "3.2 Requirements")

        if desc_anchor:
            general_desc = adf_to_text(general["fields"].get("description"))
            if general_desc:
                insert_body_text_after(desc_anchor, general_desc)

        if req_anchor:
            current = req_anchor
            greqs = sorted(
                reqs_by_uc.get(general["key"], []),
                key=lambda r: (r["fields"].get("summary", "").lower(), r["key"])
            )
            for req in greqs:
                clean_summary = clean_requirement_text(req["fields"].get("summary", ""))
                current = insert_requirement_title_after(
                    current,
                    f'{req["key"]} - {clean_summary}'
                )
                req_text = adf_to_text(req["fields"].get("description"))
                if req_text:
                    current = insert_body_text_after(current, req_text)

    # Inject Use Cases exactly after the template heading
    use_cases_anchor = find_paragraph(template, "4. Use Cases")
    image_paths = []
    current = use_cases_anchor

    for i, uc in enumerate(regular, start=1):
        if current is None:
            break

        current = insert_heading_2_after(current, f'4.{i} {uc["fields"].get("summary","")}')

        reqs = sorted(
            reqs_by_uc.get(uc["key"], []),
            key=lambda r: (r["fields"].get("summary", "").lower(), r["key"])
        )

        first = True
        for req in reqs:
            clean_summary = clean_requirement_text(req["fields"].get("summary", ""))
            current = insert_requirement_title_after(
                current,
                f'{req["key"]} - {clean_summary}'
            )

            text = adf_to_text(req["fields"].get("description"))
            if text:
                current = insert_body_text_after(current, text)

            if first:
                att = first_image_attachment(req)
                path = download_attachment(att)
                if path:
                    image_paths.append(path)
                    try:
                        img_p = insert_paragraph_after(current)
                        run = img_p.add_run()
                        run.add_picture(path, width=Inches(5.7))
                        current = img_p

                        cap = insert_paragraph_after(current)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cap.add_run(f"Figure {i} - {uc['fields'].get('summary','')}")
                        set_run_font(run, name="Arial", size=9, italic=True)
                        current = cap
                    except Exception:
                        pass
            first = False

    template.save(OUTPUT_PATH)

    for p in image_paths:
        try:
            os.unlink(p)
        except Exception:
            pass

    print(f"Saved {OUTPUT_PATH}")


def build_ssd_docx(author: str) -> str:
    main()
    return OUTPUT_PATH


if __name__ == "__main__":
    main()
